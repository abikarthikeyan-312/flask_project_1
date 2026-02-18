# app/services/question_paper_docx_service.py

import pandas as pd
from io import BytesIO
import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =========================================================
# HELPER FUNCTIONS
# =========================================================

def _safe_str(x, default=""):
    return default if x is None else str(x)

def _section_label(sec_key: str) -> str:
    return sec_key[-1].upper() if isinstance(sec_key, str) and len(sec_key) >= 4 else "?"

def _add_border_to_paragraph(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    pPr.append(pBdr)

def _add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run2 = paragraph.add_run(" of ")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)

def _setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.left_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    _add_page_number(section)

# ✅ NEW HELPER: ADDS STATUS AND TIMESTAMP
def _add_status_header(doc: Document, paper):
    """
    Adds a small metadata header at the top left with Status and Download Time.
    """
    # 1. Prepare Data
    status = paper.status.upper() if paper.status else "UNKNOWN"
    current_time = datetime.datetime.now().strftime("%d-%b-%Y %I:%M %p")
    
    # 2. Create Paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    
    # 3. Status Run (Color Coded)
    run_status = p.add_run(f"STATUS: {status}")
    run_status.bold = True
    run_status.font.name = "Courier New"
    run_status.font.size = Pt(9)
    
    if status == "ACTIVE":
        run_status.font.color.rgb = RGBColor(0, 0, 0) # Black
    elif status == "UNDER_SCRUTINY":
         run_status.font.color.rgb = RGBColor(255, 140, 0) # Orange
    else:
        run_status.font.color.rgb = RGBColor(255, 0, 0) # Red (Draft/Archived)

    # 4. Timestamp Run
    run_time = p.add_run(f"  |  DOWNLOADED: {current_time}")
    run_time.font.name = "Courier New"
    run_time.font.size = Pt(9)
    run_time.font.color.rgb = RGBColor(80, 80, 80) # Dark Grey

    # 5. Separator Line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(12)
    p_line.paragraph_format.line_spacing = Pt(6)
    run_line = p_line.add_run("_" * 95)
    run_line.font.size = Pt(6)
    run_line.font.color.rgb = RGBColor(200, 200, 200) # Light Grey

def _add_header(doc: Document, subject_obj, pattern_data: dict, semester: int):
    reg_box = doc.add_paragraph()
    reg_box.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reg_box.paragraph_format.space_after = Pt(0)
    reg_box.paragraph_format.left_indent = Inches(4.5)
    r = reg_box.add_run("REG NO :")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    _add_border_to_paragraph(reg_box)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("GURU NANAK COLLEGE (AUTONOMOUS), CHENNAI – 42.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    current_year = datetime.datetime.now().year
    exam_session = f"APRIL {current_year}" if semester % 2 == 0 else f"NOV {current_year}"
    
    exam_p = doc.add_paragraph()
    exam_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_p.paragraph_format.space_after = Pt(0)
    r = exam_p.add_run(exam_session)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    code_p = doc.add_paragraph()
    code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_p.paragraph_format.space_after = Pt(0)
    r = code_p.add_run(_safe_str(subject_obj.code))
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    marks_p = doc.add_paragraph()
    marks_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    marks_p.paragraph_format.space_after = Pt(0)
    r = marks_p.add_run(f"MAX. MARKS: {pattern_data.get('total_marks', 100)}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    time_p = doc.add_paragraph()
    time_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_p.paragraph_format.space_after = Pt(12)
    r = time_p.add_run("TIME : 3 HRS.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

def _add_section_heading(doc: Document, sec_key: str, answer_count: int, marks: int, note: str):
    label = _section_label(sec_key)
    total = answer_count * marks
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_after = Pt(0)
    head.paragraph_format.space_before = Pt(6)
    r = head.add_run(f"SECTION - {label} ({answer_count} X {marks} = {total} MARKS)")
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_p.paragraph_format.space_after = Pt(6)
    r = note_p.add_run(f"({_safe_str(note, 'Answer as required')})")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# =========================================================
# 1. DRAFT GENERATOR (Regular format)
# =========================================================
def generate_question_paper_docx(paper):
    doc = Document()
    _setup_document(doc)
    
    # ✅ INSERT STATUS HEADER
    _add_status_header(doc, paper)

    subject = paper.subject_version.subject
    semester = paper.subject_version.semester
    db_pattern = paper.subject_version.pattern.structure_json
    pattern_data = {"total_marks": paper.subject_version.pattern.total_marks}
    
    if "sections" in db_pattern:
        for key, cfg in db_pattern["sections"].items():
            pattern_data[f"Sec{key}"] = cfg.copy()

    paper_data = []
    for item in sorted(paper.items, key=lambda x: x.order_index):
        paper_data.append({"Question": item.display_text, "Marks": item.marks, "Section": item.section})

    _add_header(doc, subject, pattern_data, semester)
    df = pd.DataFrame(paper_data)
    global_qno = 1

    for sec_code in ["A", "B", "C"]:
        sec_key = f"Sec{sec_code}"
        if sec_key not in pattern_data: continue
        sec_cfg = pattern_data[sec_key]
        marks = int(sec_cfg.get("marks", 0))
        answer_count = int(sec_cfg.get("count", 0))
        note = sec_cfg.get("note", f"Answer Any {answer_count} Questions" if sec_cfg.get('total', answer_count) != answer_count else "Answer All Questions")
        
        sec_qs = df[df["Section"] == sec_code]
        if sec_qs.empty: continue

        _add_section_heading(doc, sec_key, answer_count, marks, note)
        for _, row in sec_qs.iterrows():
            p = doc.add_paragraph()
            r = p.add_run(f"{global_qno}. {row.get('Question', '')}")
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"
            global_qno += 1

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(12)
    r = footer.add_run("******")
    r.bold = True
    r.font.size = Pt(12)
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# =========================================================
# 2. OFFICIAL GENERATOR (Table format from your prompt)
# =========================================================
def generate_official_docx(paper):
    doc = Document()
    
    # ✅ INSERT STATUS HEADER
    _add_status_header(doc, paper)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    _add_page_number(section)

    subject_obj = paper.subject_version.subject
    semester = paper.subject_version.semester
    db_pattern = paper.subject_version.pattern.structure_json
    
    pattern_data = {"total_marks": paper.subject_version.pattern.total_marks}
    if "sections" in db_pattern:
        for key, cfg in db_pattern["sections"].items():
            pattern_data[f"Sec{key}"] = cfg.copy()

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("GURU NANAK COLLEGE (AUTONOMOUS), CHENNAI – 42.\n").bold = True
    current_year = datetime.datetime.now().year
    exam_session = f"APRIL {current_year}" if semester % 2 == 0 else f"NOV {current_year}"
    header.add_run(f"{exam_session}\n").bold = True
    header.add_run(f"{subject_obj.name}\n").bold = True
    header.add_run(f"{subject_obj.code}").bold = True
    
    marks_time = doc.add_paragraph()
    marks_time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    marks_time.add_run(
        f"MAX. MARKS: {pattern_data.get('total_marks', '')}\nTIME: 3 HRS."
    ).bold = True

    paper_data = []
    for item in sorted(paper.items, key=lambda x: x.order_index):
        paper_data.append({
            "Question": item.display_text,
            "Marks": item.marks,
            "Unit": item.unit,
            "Section": item.section,
            "K Level": item.k_level if item.k_level else "N/A" 
        })

    df = pd.DataFrame(paper_data)
    global_qno = 1

    col_widths = [Inches(0.6), Inches(4.3), Inches(0.7), Inches(1.0), Inches(0.9)]
    k_level_headers = {"SecA": "K1 / K2", "SecB": "K3 / K4", "SecC": "K4 / K5 / K6"}
    
    for sec_code in ["A", "B", "C"]:
        sec_key = f"Sec{sec_code}"
        if sec_key not in pattern_data: 
            continue
            
        sec_cfg = pattern_data[sec_key]        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False 

        hdr_cells = table.rows[0].cells
        for i, width in enumerate(col_widths):
            hdr_cells[i].width = width

        hdr_cells[0].text = "Q. No"
        total_sec_marks = int(sec_cfg['count']) * int(sec_cfg['marks'])
        instruction = f"SECTION - {sec_code} " \
                      f"({sec_cfg['count']} X {sec_cfg['marks']} = {total_sec_marks} MARKS)\n" \
                      f"{sec_cfg.get('note', 'Answer as required')}"
        hdr_cells[1].text = instruction
        hdr_cells[2].text = "Marks"
        hdr_cells[3].text = "Course Outcome"
        k_range = k_level_headers.get(sec_key, "K-Level")
        hdr_cells[4].text = f"K-Level\n{k_range}"        
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                for run in paragraph.runs:
                    run.bold = True

        sec_qs = df[df["Section"] == sec_code]

        for _, row in sec_qs.iterrows():
            row_cells = table.add_row().cells
            for i, width in enumerate(col_widths):
                row_cells[i].width = width
            
            row_cells[0].text = str(global_qno)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = str(row['Question'])
            row_cells[2].text = str(row['Marks'])
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].text = f"CO{row.get('Unit', 'N/A')}" 
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].text = f"{row.get('K Level', 'N/A')}"
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            global_qno += 1

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream